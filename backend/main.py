from fastapi import FastAPI, Form, UploadFile, File, Request, HTTPException
from fastapi.responses import FileResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from typing import List, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import uuid
import re
from datetime import datetime, timedelta
import io
from bs4 import BeautifulSoup
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

app = FastAPI()

# Directory setup
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(BASE_DIR, "output")
STATIC_DIR = os.path.join(BASE_DIR, "static")
UPLOAD_DIR = os.path.join(BASE_DIR, "uploads")
TEMPLATES_DIR = os.path.join(BASE_DIR, "templates")

os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(STATIC_DIR, exist_ok=True)
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(TEMPLATES_DIR, exist_ok=True)

app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")
templates = Jinja2Templates(directory=TEMPLATES_DIR)

def cleanup_old_files(directory, hours=24):
    now = datetime.now()
    for filename in os.listdir(directory):
        file_path = os.path.join(directory, filename)
        if os.path.isfile(file_path):
            file_time = datetime.fromtimestamp(os.path.getmtime(file_path))
            if now - file_time > timedelta(hours=hours):
                try:
                    os.remove(file_path)
                except Exception as e:
                    print(f"Failed to remove {file_path}: {e}")

@app.on_event("startup")
async def startup_event():
    cleanup_old_files(OUTPUT_DIR)

@app.get("/", response_class=HTMLResponse)
async def read_root(request: Request):
    try:
        return templates.TemplateResponse("index.html", {"request": request})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Template error: {str(e)}")

@app.post("/preview_report")
async def preview_report(
    request: Request,
    eventType: str = Form(...),
    department: str = Form(...),
    topic: str = Form(...),
    expertName: Optional[str] = Form(None),
    venue: str = Form(...),
    eventDurationType: str = Form(...),
    date: Optional[str] = Form(None),
    startTime: Optional[str] = Form(None),
    endTime: Optional[str] = Form(None),
    startDate: Optional[str] = Form(None),
    endDate: Optional[str] = Form(None),
    coordinator: str = Form(...),
    participants: int = Form(...),
    summary: str = Form(...),
    outcome: str = Form(...),
    hodName: str = Form(...)
):
    try:
        if eventDurationType.lower() == "multiple" and startDate and endDate:
            formatted_dateTime = (
                datetime.strptime(startDate, "%Y-%m-%d").strftime("%B %d, %Y")
                + " to " +
                datetime.strptime(endDate, "%Y-%m-%d").strftime("%B %d, %Y")
            )
        else:
            formatted_date = datetime.strptime(date, "%Y-%m-%d").strftime("%B %d, %Y")
            formatted_time = (
                datetime.strptime(startTime, "%H:%M").strftime("%I:%M %p")
                + " - " +
                datetime.strptime(endTime, "%H:%M").strftime("%I:%M %p")
            )
            formatted_dateTime = formatted_date + ", " + formatted_time

        report_data = {
            "eventType": eventType.title(),
            "department": department,
            "topic": topic,
            "expertName": expertName if expertName else "N/A",
            "venue": venue,
            "dateTime": formatted_dateTime,
            "coordinator": coordinator,
            "participants": participants,
            "summary": summary,
            "outcome": outcome,
            "hodName": hodName
        }
        return templates.TemplateResponse("preview.html", {"request": request, "report": report_data})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Preview generation failed: {str(e)}")

def replace_placeholder(doc, placeholder, value):
    def process_paragraph(paragraph, ph, val):
        if ph in paragraph.text:
            full_text = ''.join(run.text for run in paragraph.runs)
            if ph in full_text:
                paragraph.clear()
                new_run = paragraph.add_run(full_text.replace(ph, str(val)))
                new_run.font.name = 'DIN Pro Regular'
                new_run.font.size = Pt(11)
                # Set color to dark blue for specific text
                if paragraph.text.strip().endswith(" Report") or paragraph.text.strip().startswith("Department of "):
                    new_run.font.color.rgb = RGBColor(0, 0, 139)
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph, placeholder, value)
    def process_table(table):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph, placeholder, value)
                for nested_table in cell.tables:
                    process_table(nested_table)
    for table in doc.tables:
        process_table(table)

def process_node_formatting(paragraph, node):
    if isinstance(node, str):
        run = paragraph.add_run(node)
        run.font.name = 'DIN Pro Regular'
        run.font.size = Pt(11)
        return run
    if not hasattr(node, 'name'):
        run = paragraph.add_run(str(node))
        run.font.name = 'DIN Pro Regular'
        run.font.size = Pt(11)
        return run
    if node.name in ['strong', 'b']:
        last_run = None
        if node.contents:
            for child in node.contents:
                last_run = process_node_formatting(paragraph, child)
                if last_run is not None:
                    last_run.bold = True
        else:
            last_run = paragraph.add_run(node.get_text())
            last_run.bold = True
            last_run.font.name = 'DIN Pro Regular'
            last_run.font.size = Pt(11)
        return last_run
    elif node.name in ['em', 'i']:
        last_run = None
        if node.contents:
            for child in node.contents:
                last_run = process_node_formatting(paragraph, child)
                if last_run is not None:
                    last_run.italic = True
        else:
            last_run = paragraph.add_run(node.get_text())
            last_run.italic = True
            last_run.font.name = 'DIN Pro Regular'
            last_run.font.size = Pt(11)
        return last_run
    elif node.contents:
        last_run = None
        for child in node.contents:
            last_run = process_node_formatting(paragraph, child)
        return last_run
    else:
        run = paragraph.add_run(node.get_text())
        run.font.name = 'DIN Pro Regular'
        run.font.size = Pt(11)
        return run

def insert_paragraph_after(paragraph, text='', style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style:
        new_para.style = style
    return new_para

def replace_placeholder_with_html(doc, placeholder, html_content):
    paragraphs = list(doc.paragraphs)
    for i, paragraph in enumerate(paragraphs):
        if placeholder in paragraph.text:
            paragraph.clear()
            soup = BeautifulSoup(html_content, "html.parser")
            last_par = paragraph
            elements = list(soup.body.children) if soup.body else list(soup.children)
            if not elements:
                last_par.add_run(soup.get_text())
                return
            first = True
            for element in elements:
                if isinstance(element, str) and element.strip() == '':
                    continue
                if element.name == 'p':
                    if first:
                        new_par = last_par
                        first = False
                    else:
                        new_par = insert_paragraph_after(last_par)
                    for child in element.children:
                        process_node_formatting(new_par, child)
                    last_par = new_par
                elif element.name in ['ul', 'ol']:
                    for li in element.find_all('li', recursive=False):
                        list_style = 'List Bullet' if element.name == 'ul' else 'List Number'
                        new_par = insert_paragraph_after(last_par, style=list_style)
                        for child in li.children:
                            process_node_formatting(new_par, child)
                        last_par = new_par
                elif isinstance(element, str) and element.strip():
                    new_par = last_par if first else insert_paragraph_after(last_par)
                    run = new_par.add_run(element.strip())
                    run.font.name = 'DIN Pro Regular'
                    run.font.size = Pt(11)
                    last_par = new_par
            break

def update_header(doc, eventType):
    for paragraph in doc.paragraphs:
        if "{{eventType}}" in paragraph.text:
            text = paragraph.text.replace("{{eventType}}", eventType.title())
            paragraph.clear()
            run = paragraph.add_run(text)
            run.font.name = 'DIN Pro Regular'
            run.font.size = Pt(24)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 139)  # Dark blue
            break

def set_section_vertical_alignment_bottom(section):
    try:
        sectPr = section._sectPr
        for child in sectPr.findall(qn('w:valign')):
            sectPr.remove(child)
        vAlign = OxmlElement('w:valign')
        vAlign.set(qn('w:val'), 'bottom')
        sectPr.append(vAlign)
    except Exception as e:
        print(f"Could not set vertical alignment: {e}")

def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is not None:
        tblPr.remove(tblBorders)

@app.post("/generate_report")
async def generate_report(
    eventType: str = Form(...),
    department: str = Form(...),
    topic: str = Form(...),
    expertName: Optional[str] = Form(None),
    venue: str = Form(...),
    eventDurationType: str = Form(...),
    date: Optional[str] = Form(None),
    startTime: Optional[str] = Form(None),
    endTime: Optional[str] = Form(None),
    startDate: Optional[str] = Form(None),
    endDate: Optional[str] = Form(None),
    coordinator: str = Form(...),
    participants: int = Form(...),
    summary: str = Form(...),
    outcome: str = Form(...),
    hodName: str = Form(...),
    invitePoster: Optional[List[UploadFile]] = File(None),
    actionPhotos: Optional[List[UploadFile]] = File(None),
    attendanceSheet: Optional[List[UploadFile]] = File(None),
    analysisReport: Optional[List[UploadFile]] = File(None)
):
    try:
        if eventDurationType.lower() == "multiple" and startDate and endDate:
            formatted_dateTime = (
                datetime.strptime(startDate, "%Y-%m-%d").strftime("%B %d, %Y")
                + " to " +
                datetime.strptime(endDate, "%Y-%m-%d").strftime("%B %d, %Y")
            )
        else:
            formatted_date = datetime.strptime(date, "%Y-%m-%d").strftime("%B %d, %Y")
            formatted_time = (
                datetime.strptime(startTime, "%H:%M").strftime("%I:%M %p")
                + " - " +
                datetime.strptime(endTime, "%H:%M").strftime("%I:%M %p")
            )
            formatted_dateTime = formatted_date + ", " + formatted_time

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_topic = re.sub(r'[^\w\s]', '_', topic)
        output_filename = f"{department}_{safe_topic.replace(' ', '_')}_{timestamp}.docx"
        output_path = os.path.join(OUTPUT_DIR, output_filename)

        template_name = f"{eventType.lower()}_template.docx"
        template_path = os.path.join(TEMPLATES_DIR, template_name)
        if not os.path.exists(template_path):
            template_path = os.path.join(TEMPLATES_DIR, "workshop_template.docx")
            if not os.path.exists(template_path):
                doc = Document()
                doc.add_paragraph("{{eventType}} Report")
            else:
                doc = Document(template_path)
        else:
            doc = Document(template_path)

        replacements = {
            "{{eventType}}": eventType.title(),
            "{{department}}": department,
            "{{topic}}": topic,
            "{{expertName}}": expertName if expertName else "N/A",
            "{{venue}}": venue,
            "{{dateTime}}": formatted_dateTime,
            "{{coordinator}}": coordinator,
            "{{participants}}": str(participants),
            "{{hodName}}": hodName
        }
        for placeholder, value in replacements.items():
            replace_placeholder(doc, placeholder, value)
        replace_placeholder_with_html(doc, "{{summary}}", summary)
        replace_placeholder_with_html(doc, "{{outcome}}", outcome)
        update_header(doc, eventType)

        image_sections = [
            ("Invite Poster", invitePoster),
            ("Action Photos", actionPhotos),
            ("Attendance Sheet", attendanceSheet),
            ("Analysis Report", analysisReport)
        ]
        for section_name, images in image_sections:
            valid_images = [img for img in images if img and img.filename] if images else []
            if valid_images:
                doc.add_page_break()
                p = doc.add_paragraph()
                run = p.add_run(section_name)
                run.bold = True
                run.font.size = Pt(16)
                run.font.name = 'DIN Pro Regular'
                for img in valid_images:
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    unique_filename = f"{uuid.uuid4()}_{img.filename}"
                    img_path = os.path.join(UPLOAD_DIR, unique_filename)
                    with open(img_path, "wb") as buffer:
                        await img.seek(0)
                        contents = await img.read()
                        buffer.write(contents)
                    try:
                        p_img.add_run().add_picture(img_path, width=Inches(6))
                        os.remove(img_path)
                    except Exception as e:
                        print(f"Error processing image {img.filename}: {e}")

        # Add signature table in a new section aligned to bottom
        doc.add_section()
        set_section_vertical_alignment_bottom(doc.sections[-1])
        sig_table = doc.add_table(rows=1, cols=2)
        sig_table.style = None  # No default style to ensure borders are controlled
        remove_table_borders(sig_table)
        sig_table.cell(0, 0).text = "Name & Signature of Faculty-in-charge\n" + coordinator
        sig_table.cell(0, 1).text = "Name & Signature of HoD\n" + hodName
        for row in sig_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.name = 'DIN Pro Regular'
                        run.font.size = Pt(11)

        doc.save(output_path)
        cleanup_old_files(OUTPUT_DIR, hours=24)
        return {
            "message": "Report generated successfully",
            "filename": output_filename,
            "download_url": f"/download_report/{output_filename}"
        }
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"error": f"Failed to generate report: {str(e)}"}

@app.get("/download_report/{filename}")
async def download_report(filename: str):
    sanitized_filename = os.path.basename(filename)
    if sanitized_filename != filename:
        raise HTTPException(status_code=400, detail="Invalid filename")
    file_path = os.path.join(OUTPUT_DIR, sanitized_filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(
        file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=sanitized_filename
    )

@app.on_event("startup")
async def create_default_template():
    default_template_path = os.path.join(TEMPLATES_DIR, "workshop_template.docx")
    if not os.path.exists(default_template_path):
        try:
            doc = Document()
            p_title = doc.add_paragraph()
            run_title = p_title.add_run("{{eventType}} Report")
            run_title.font.name = 'DIN Pro Regular'
            run_title.font.size = Pt(24)
            run_title.bold = True

            p_details = doc.add_paragraph("Event Details")
            p_details.runs[0].font.name = 'DIN Pro Regular'
            p_details.runs[0].font.size = Pt(16)
            p_details.runs[0].bold = True

            details = [
                ("Department", "{{department}}"),
                ("Topic", "{{topic}}"),
                ("Expert Name", "{{expertName}}"),
                ("Venue", "{{venue}}"),
                ("Event Date/Time", "{{dateTime}}"),
                ("Faculty Coordinator", "{{coordinator}}"),
                ("HOD Name", "{{hodName}}"),
                ("Participants", "{{participants}}")
            ]
            table = doc.add_table(rows=len(details), cols=2)
            table.style = 'Table Grid'
            for i, (label, value) in enumerate(details):
                row = table.rows[i]
                row.cells[0].text = label
                row.cells[1].text = value
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'DIN Pro Regular'
                            run.font.size = Pt(11)
            p_summary = doc.add_paragraph("Summary")
            p_summary.runs[0].font.name = 'DIN Pro Regular'
            p_summary.runs[0].font.size = Pt(16)
            p_outcome = doc.add_paragraph("Outcome")
            p_outcome.runs[0].font.name = 'DIN Pro Regular'
            p_outcome.runs[0].font.size = Pt(16)
            # Signature table is added dynamically in generate_report, not here
            doc.save(default_template_path)
            print(f"Created default template at {default_template_path}")
        except Exception as e:
            print(f"Failed to create default template: {e}")

@app.on_event("startup")
async def schedule_cleanup():
    import asyncio
    async def periodic_cleanup():
        while True:
            cleanup_old_files(OUTPUT_DIR)
            await asyncio.sleep(3600)
    asyncio.create_task(periodic_cleanup())

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)